from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_paragraph()
title_run = title.add_run('StayNGo DevOps & Infrastructure Documentation')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('Complete DevOps Architecture, CI/CD Pipeline & Deployment Strategy')
subtitle_run.font.size = Pt(14)
subtitle_run.font.italic = True
subtitle_run.font.color.rgb = RGBColor(64, 64, 64)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

# Executive Summary
doc.add_heading('Executive Summary', 1)
summary_text = """StayNGo implements a comprehensive, enterprise-grade DevOps infrastructure designed for scalability, reliability, and rapid deployment cycles. This documentation provides an in-depth overview of all DevOps technologies, their integration points, and operational procedures. The infrastructure is built on AWS cloud computing, automated through Jenkins CI/CD pipelines, containerized with Docker, and orchestrated using both Docker Compose and Kubernetes."""
doc.add_paragraph(summary_text)

# Table of Contents
doc.add_heading('Table of Contents', 1)
toc_items = [
    '1. Overview of DevOps Architecture',
    '2. Version Control System (GitHub)',
    '3. Continuous Integration/Continuous Deployment (Jenkins)',
    '4. Code Quality Analysis (SonarQube)',
    '5. Containerization (Docker & Docker Compose)',
    '6. Infrastructure as Code (Terraform)',
    '7. Configuration Management (Ansible)',
    '8. Container Orchestration (Kubernetes)',
    '9. Cloud Infrastructure (AWS)',
    '10. Monitoring & Logging',
    '11. Security & Best Practices',
    '12. Deployment Pipeline Workflow',
    '13. Troubleshooting & Maintenance'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Section 1: Overview
doc.add_heading('1. Overview of DevOps Architecture', 1)

doc.add_heading('1.1 Architecture Diagram', 2)
doc.add_paragraph('[IMAGE PLACEHOLDER: DevOps Architecture Diagram]')
doc.add_paragraph('Insert a comprehensive architecture diagram showing:')
doc.add_paragraph('• GitHub at the source', style='List Bullet')
doc.add_paragraph('• Jenkins pipeline flow', style='List Bullet')
doc.add_paragraph('• SonarQube integration', style='List Bullet')
doc.add_paragraph('• Docker build and push', style='List Bullet')
doc.add_paragraph('• AWS EC2 deployment targets', style='List Bullet')

doc.add_heading('1.2 Technology Stack Overview', 2)
tech_data = [
    ('Technology', 'Purpose', 'Version/Type'),
    ('Jenkins', 'CI/CD Orchestration & Automation', 'Latest LTS'),
    ('SonarQube', 'Code Quality & Static Analysis', 'Community Edition'),
    ('Docker', 'Container Runtime & Build System', 'Latest'),
    ('Docker Compose', 'Multi-container Orchestration', 'V2'),
    ('Kubernetes (k3s)', 'Container Orchestration (Optional)', 'Latest'),
    ('Terraform', 'Infrastructure as Code (IaC)', '~> 5.0'),
    ('Ansible', 'Configuration Management & Provisioning', 'Latest'),
    ('AWS EC2', 'Cloud Compute Resources', 'Ubuntu 22.04 LTS'),
    ('GitHub', 'Version Control & Webhooks', 'Cloud-based'),
    ('Docker Hub', 'Container Image Registry', 'Public Registry'),
    ('Supabase', 'Database & Authentication', 'PostgreSQL Backend')
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Technology'
hdr_cells[1].text = 'Purpose'
hdr_cells[2].text = 'Version/Type'

for row_data in tech_data[1:]:
    row_cells = table.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]

doc.add_page_break()

# Section 2: GitHub
doc.add_heading('2. Version Control System (GitHub)', 1)

doc.add_heading('2.1 What is GitHub?', 2)
github_desc = """GitHub is a cloud-based Git repository hosting service that provides version control, collaboration features, and integration points for automated workflows. It serves as the single source of truth for all source code and enables webhook-based triggering of CI/CD pipelines."""
doc.add_paragraph(github_desc)

doc.add_heading('2.2 Key Features in StayNGo', 2)
doc.add_paragraph('Repository Structure:', style='List Bullet')
doc.add_paragraph('• Main branch: Production-ready code', style='List Bullet 2')
doc.add_paragraph('• Development branches: Feature development', style='List Bullet 2')
doc.add_paragraph('• Pull requests: Code review mechanism', style='List Bullet 2')
doc.add_paragraph('• Webhook integration: Automatic Jenkins triggers', style='List Bullet')
doc.add_paragraph('• Branch protection rules: Enforce quality gates', style='List Bullet')

doc.add_heading('2.3 Webhook Configuration', 2)
doc.add_paragraph('[IMAGE PLACEHOLDER: GitHub Webhook Configuration]')
webhook_config = """URL: Jenkins Server IP:8080/github-webhook/
Events: Push events
Authentication: Automatically verified by Jenkins
This webhook automatically triggers the Jenkins pipeline whenever code is pushed to the main branch."""
doc.add_paragraph(webhook_config)

doc.add_page_break()

# Section 3: Jenkins
doc.add_heading('3. Continuous Integration/Continuous Deployment (Jenkins)', 1)

doc.add_heading('3.1 What is Jenkins?', 2)
jenkins_desc = """Jenkins is an open-source automation server that orchestrates complex build, test, and deployment workflows. It provides a robust framework for Continuous Integration (CI) and Continuous Deployment (CD) through declarative pipelines, extensive plugin ecosystem, and webhook-based triggering mechanisms."""
doc.add_paragraph(jenkins_desc)

doc.add_heading('3.2 Jenkins Architecture in StayNGo', 2)
doc.add_paragraph('[IMAGE PLACEHOLDER: Jenkins Pipeline Flow]')

doc.add_heading('3.3 Pipeline Stages Explained', 2)

doc.add_heading('Stage 1: Checkout', 3)
doc.add_paragraph('Pulls the latest code from GitHub repository. Uses Git plugin to clone the repository into Jenkins workspace.')

doc.add_heading('Stage 2: SonarQube Code Quality Check', 3)
doc.add_paragraph('Executes SonarQube Scanner to analyze code for bugs, vulnerabilities, code smells, and coverage metrics. Enforces quality gates - if standards are not met, pipeline fails.')

doc.add_heading('Stage 3: Build Docker Images', 3)
doc.add_paragraph('Builds Docker images for both frontend and backend using multi-stage Dockerfile. Uses --no-cache flag to ensure latest dependencies are pulled. Injects environment variables at build time.')

doc.add_heading('Stage 4: Push to Registry', 3)
doc.add_paragraph('Authenticates with Docker Hub using stored credentials. Pushes built images to Docker Hub registry: mayankc1533262/stayngo-backend:latest and mayankc1533262/stayngo-frontend:latest')

doc.add_heading('Stage 5: Deploy to AWS EC2', 3)
doc.add_paragraph('Establishes SSH connection to EC2 instance using private key authentication. Pulls latest images from Docker Hub. Executes docker-compose up -d to deploy containers on production.')

doc.add_heading('3.4 Jenkins Configuration Details', 2)
config_details = """Dashboard URL: http://65.0.117.219:8080

Credentials Stored:
  • Docker Hub username/password
  • SonarQube authentication token
  • AWS EC2 SSH private key
  • GitHub personal access token (optional)

Plugin Requirements:
  • Docker plugin
  • SonarQube Scanner plugin
  • GitHub plugin
  • SSH plugin
  • Email plugin (for notifications)"""
doc.add_paragraph(config_details)

doc.add_page_break()

# Section 4: SonarQube
doc.add_heading('4. Code Quality Analysis (SonarQube)', 1)

doc.add_heading('4.1 What is SonarQube?', 2)
sonar_desc = """SonarQube is an open-source platform dedicated to continuous inspection of code quality. It performs static code analysis to identify bugs, vulnerabilities, code smells, and coverage issues. SonarQube provides detailed reports and enforces quality gates that can block builds if standards are not met."""
doc.add_paragraph(sonar_desc)

doc.add_heading('4.2 Key Metrics Tracked', 2)
doc.add_paragraph('[IMAGE PLACEHOLDER: SonarQube Metrics Dashboard]')
metrics = [
    'Code Coverage: Percentage of code covered by tests',
    'Bugs: Number of confirmed bugs detected',
    'Vulnerabilities: Security issues and exploitable weaknesses',
    'Code Smells: Code patterns that violate best practices',
    'Duplication: Percentage of duplicated code',
    'Technical Debt: Estimated effort to fix all issues',
    'Security Hotspots: Areas requiring security review'
]
for metric in metrics:
    doc.add_paragraph(metric, style='List Bullet')

doc.add_heading('4.3 Quality Gates', 2)
qg_text = """Quality Gates are predefined thresholds that determine whether a code analysis should be marked as passed or failed. 
In StayNGo, the following gates are enforced:

  • Code Coverage: Minimum 30%
  • New Bugs: Maximum 0
  • New Vulnerabilities: Maximum 0
  • Duplicated Lines: Less than 5%
  • Maintainability Rating: A or B

If any gate fails, the Jenkins pipeline stops and deployment is prevented."""
doc.add_paragraph(qg_text)

doc.add_heading('4.4 Dashboard & Reports', 2)
doc.add_paragraph('Access URL: http://65.0.117.219:9000')
doc.add_paragraph('Features:', style='List Bullet')
doc.add_paragraph('• Real-time code analysis results', style='List Bullet 2')
doc.add_paragraph('• Historical trend reports', style='List Bullet 2')
doc.add_paragraph('• Issue tracking and prioritization', style='List Bullet 2')
doc.add_paragraph('• Integration with IDE plugins', style='List Bullet 2')

doc.add_page_break()

# Section 5: Docker
doc.add_heading('5. Containerization (Docker & Docker Compose)', 1)

doc.add_heading('5.1 What is Docker?', 2)
docker_desc = """Docker is a containerization platform that packages applications and their dependencies into isolated, lightweight, portable containers. Each container includes the application code, runtime, libraries, and environment variables needed to run the application independently across different environments."""
doc.add_paragraph(docker_desc)

doc.add_heading('5.2 Container Benefits', 2)
doc.add_paragraph('[IMAGE PLACEHOLDER: Docker Architecture Diagram]')
benefits = [
    'Consistency: Same container runs identically on laptop, CI/CD server, and production',
    'Isolation: Applications do not interfere with each other',
    'Efficiency: Lightweight compared to virtual machines',
    'Scalability: Easy to spin up multiple instances',
    'Portability: Runs on any system with Docker installed'
]
for benefit in benefits:
    doc.add_paragraph(benefit, style='List Bullet')

doc.add_heading('5.3 Frontend Docker Image', 2)
frontend_info = """Base Image: node:20-alpine
Application: Next.js React application

Build Process:
  1. Install npm dependencies from package.json
  2. Build Next.js application (npm run build)
  3. Inject NEXT_PUBLIC_* environment variables at build time
  4. Expose port 3000
  5. Run: npm start

Environment Variables Required:
  • NEXT_PUBLIC_SUPABASE_URL
  • NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY
  • NEXT_PUBLIC_API_URL"""
doc.add_paragraph(frontend_info)

doc.add_heading('5.4 Backend Docker Image', 2)
backend_info = """Base Image: python:3.11-slim
Application: Flask API with Gunicorn WSGI server

Build Process:
  1. Install system dependencies
  2. Install Python dependencies from requirements.txt
  3. Set up Flask application
  4. Expose port 5001
  5. Run: gunicorn --workers 3 --threads 4 app:app

Gunicorn Configuration:
  • Workers: 3 (multi-process concurrency)
  • Threads: 4 (per-worker threading for I/O operations)
  • Bind: 0.0.0.0:5001 (listen on all interfaces)
  • Timeout: 120 seconds (prevents DDoS from slow clients)"""
doc.add_paragraph(backend_info)

doc.add_heading('5.5 Docker Compose V2', 2)
compose_desc = """Docker Compose is a tool for defining and running multi-container Docker applications. A YAML file (docker-compose.prod.yml) defines all services, their configurations, and networking.

Services in StayNGo:
  1. Frontend Service
     - Image: mayankc1533262/stayngo-frontend:latest
     - Port: 3000:3000
     - Restart: always
     - Environment: Loaded from .env file

  2. Backend Service
     - Image: mayankc1533262/stayngo-backend:latest
     - Port: 5001:5001
     - Restart: always
     - Environment: Loaded from .env file

Commands:
  • docker-compose up -d: Start services in background
  • docker-compose down: Stop and remove services
  • docker-compose logs -f: View live logs
  • docker-compose pull: Update images"""
doc.add_paragraph(compose_desc)

# Save part 1
doc.save(r'c:\Users\mayan\OneDrive\Desktop\Mayank\StayNGo\StayNGo_DevOps_Documentation_Part1.docx')
print('Part 1 created successfully')

from docx import Document

# Load part 1 document
doc = Document(r'c:\Users\mayan\OneDrive\Desktop\Mayank\StayNGo\StayNGo_DevOps_Documentation_Part1.docx')

doc.add_page_break()

# Section 6: Terraform
doc.add_heading('6. Infrastructure as Code (Terraform)', 1)

doc.add_heading('6.1 What is Terraform?', 2)
terraform_desc = """Terraform is an Infrastructure as Code (IaC) tool that allows defining cloud infrastructure in declarative configuration files. Instead of manually clicking through AWS console, Terraform automates infrastructure provisioning, making it reproducible, version-controlled, and maintainable."""
doc.add_paragraph(terraform_desc)

doc.add_heading('6.2 Infrastructure Components Defined', 2)
doc.add_paragraph('[IMAGE PLACEHOLDER: AWS Infrastructure Diagram]')
components = [
    'AWS VPC (Virtual Private Cloud): Network isolation',
    'Security Groups: Firewall rules controlling inbound/outbound traffic',
    'EC2 Instance: Virtual machine for running applications',
    'EBS Volume: Block storage (20GB gp3) for persistent data',
    'SSH Key Pair: Public/private key for secure access',
    'IAM Roles: (Optional) for AWS service permissions'
]
for comp in components:
    doc.add_paragraph(comp, style='List Bullet')

doc.add_heading('6.3 Key Terraform Files', 2)
tf_files = """main.tf: Main infrastructure definition
  • Defines AWS provider region (ap-south-1)
  • Creates security group with port rules
  • Creates EC2 instance with specified configuration
  • Outputs instance public IP

variables.tf: Variable definitions
  • aws_region: AWS region (default: ap-south-1)
  • instance_type: EC2 instance type (default: m7i-flex.large)
  • ami_id: Ubuntu AMI ID to use

terraform.tfvars: Variable values (not tracked in Git)
  • Contains specific values for variables
  • Secrets and sensitive data"""
doc.add_paragraph(tf_files)

doc.add_heading('6.4 Terraform Workflow', 2)
workflow = """1. terraform init: Initialize Terraform working directory, download providers
2. terraform plan: Review what changes will be made (dry-run)
3. terraform apply: Execute the plan and create/update infrastructure
4. terraform destroy: Tear down all infrastructure (use with caution)
5. terraform state: Maintains state file tracking actual infrastructure

Important: State file must be backed up and secured!"""
doc.add_paragraph(workflow)

doc.add_page_break()

# Section 7: Ansible
doc.add_heading('7. Configuration Management (Ansible)', 1)

doc.add_heading('7.1 What is Ansible?', 2)
ansible_desc = """Ansible is an agentless configuration management tool that automates infrastructure provisioning, configuration, and deployment. Unlike tools requiring agents on target machines, Ansible uses SSH to execute commands, making it simpler and more lightweight."""
doc.add_paragraph(ansible_desc)

doc.add_heading('7.2 Playbook Structure', 2)
doc.add_paragraph('[IMAGE PLACEHOLDER: Ansible Playbook Flow]')
playbook_structure = """File: ansible/setup.yml
Purpose: Configure EC2 instance after Terraform creates it

Key Tasks:
1. Update System Packages
   - Run apt-get update and upgrade
   - Ensure system is current

2. Install Docker
   - Install docker.io package
   - Install docker-compose
   - Install supporting tools (git, curl, htop)

3. Docker Service Setup
   - Add ubuntu user to docker group (run docker without sudo)
   - Enable Docker service to start on boot

4. Memory Configuration
   - Create 2GB swap file
   - Configure swappiness for optimal performance
   - Set vm.max_map_count for SonarQube

5. Application Setup
   - Create /home/ubuntu/stayngo directory
   - Copy docker-compose.prod.yml file
   - Set proper permissions

6. Service Startup
   - Start SonarQube container
   - Start Jenkins container with Docker socket mounted"""
doc.add_paragraph(playbook_structure)

doc.add_heading('7.3 Inventory File', 2)
doc.add_paragraph('File: ansible/inventory.ini')
doc.add_paragraph('Defines target hosts for playbook execution:')
doc.add_paragraph('[ec2]')
doc.add_paragraph('65.0.117.219 ansible_user=ubuntu ansible_ssh_private_key_file=~/.ssh/stayngo-deployer-key')

doc.add_page_break()

# Section 8: Kubernetes
doc.add_heading('8. Container Orchestration (Kubernetes)', 1)

doc.add_heading('8.1 What is Kubernetes?', 2)
k8s_desc = """Kubernetes (k8s) is an open-source container orchestration platform that automates deployment, scaling, and management of containerized applications. While StayNGo currently uses Docker Compose, Kubernetes provides advanced features like auto-scaling, self-healing, and rolling updates."""
doc.add_paragraph(k8s_desc)

doc.add_heading('8.2 Kubernetes Concepts', 2)
doc.add_paragraph('[IMAGE PLACEHOLDER: Kubernetes Architecture]')
k8s_concepts = [
    'Pod: Smallest deployable unit, contains one or more containers',
    'Deployment: Manages replicas of pods, ensures desired state',
    'Service: Exposes pods on network, provides load balancing',
    'ConfigMap: Store configuration data for applications',
    'Secret: Store sensitive data (API keys, passwords)',
    'Namespace: Virtual cluster for resource isolation'
]
for concept in k8s_concepts:
    doc.add_paragraph(concept, style='List Bullet')

doc.add_heading('8.3 Frontend Deployment (k8s/frontend.yaml)', 2)
frontend_k8s = """Deployment: stayngo-frontend
  • Replicas: 2 (high availability)
  • Image: stayngo-frontend:latest
  • Port: 3000
  • Config: Loaded from ConfigMap

Service: stayngo-frontend
  • Type: LoadBalancer (exposes on port 80)
  • Maps port 80 to container port 3000
  • Provides external access point"""
doc.add_paragraph(frontend_k8s)

doc.add_heading('8.4 Backend Deployment (k8s/backend.yaml)', 2)
backend_k8s = """Deployment: stayngo-backend
  • Replicas: 2 (high availability)
  • Image: stayngo-backend:latest
  • Port: 5001
  • Config: Loaded from ConfigMap and Secrets

Service: stayngo-backend
  • Type: ClusterIP (internal-only access)
  • Only accessible within cluster
  • Used by frontend pods to communicate"""
doc.add_paragraph(backend_k8s)

doc.add_page_break()

# Section 9: AWS
doc.add_heading('9. Cloud Infrastructure (AWS)', 1)

doc.add_heading('9.1 What is AWS?', 2)
aws_desc = """Amazon Web Services (AWS) is a comprehensive cloud computing platform providing compute, storage, networking, databases, and many other services. StayNGo leverages AWS EC2 for hosting the application infrastructure."""
doc.add_paragraph(aws_desc)

doc.add_heading('9.2 EC2 Instance Details', 2)
doc.add_paragraph('[IMAGE PLACEHOLDER: AWS EC2 Dashboard]')
ec2_details = """Instance Configuration:
  • Instance ID: Managed by Terraform
  • Instance Type: m7i-flex.large (configurable)
  • vCPU: 2 cores
  • Memory: 8 GB RAM
  • Storage: 20 GB gp3 EBS volume
  • AMI: Ubuntu 22.04 LTS
  • Region: ap-south-1 (Mumbai)
  • Public IP: 65.0.117.219 (Elastic IP)

Performance Characteristics:
  • Good for: Development, staging, small to medium production workloads
  • vCPU Credit System: Flexible performance scaling
  • Network: High bandwidth for data transfer"""
doc.add_paragraph(ec2_details)

doc.add_heading('9.3 Security Groups', 2)
sg_details = """Security Group: stayngo-compute-sg

Inbound Rules (Allow):
  • Port 22 (SSH): From 0.0.0.0/0 - Administration access
  • Port 80 (HTTP): From 0.0.0.0/0 - Frontend through load balancer
  • Port 3000 (Frontend App): From 0.0.0.0/0 - Direct access (development)
  • Port 5001 (Backend API): From 0.0.0.0/0 - Direct API access
  • Port 8080 (Jenkins): From 0.0.0.0/0 - CI/CD dashboard
  • Port 9000 (SonarQube): From 0.0.0.0/0 - Code quality dashboard

Outbound Rules (Allow):
  • All traffic to 0.0.0.0/0 - Download images, dependencies, etc.

IMPORTANT: In production, restrict SSH to specific IPs!"""
doc.add_paragraph(sg_details)

doc.add_heading('9.4 Elastic Block Store (EBS)', 2)
ebs_info = """Volume: 20 GB gp3
  • Volume Type: gp3 (General Purpose, latest generation)
  • IOPS: 3,000 baseline (scalable)
  • Throughput: 125 MB/s baseline (scalable)
  • Performance: Suitable for web applications and databases

Uses:
  • Root filesystem for OS and system files
  • Docker image storage
  • Container data persistence
  • Application logs storage

Cost: ~$2/month per 100 GB"""
doc.add_paragraph(ebs_info)

doc.add_page_break()

# Section 10: Monitoring
doc.add_heading('10. Monitoring & Logging', 1)

doc.add_heading('10.1 Application Logs', 2)
doc.add_paragraph('[IMAGE PLACEHOLDER: Log Aggregation Dashboard]')
logs_info = """Docker Logs:
  docker-compose -f docker-compose.prod.yml logs -f [service-name]
  • -f flag: Follow logs in real-time
  • View stderr and stdout from containers

System Logs:
  /var/log/syslog - System events
  /var/log/docker.log - Docker daemon logs
  journalctl - Systemd journal for service logs

Application-Specific:
  Frontend: Next.js build logs, runtime errors
  Backend: Flask application logs, API request logs"""
doc.add_paragraph(logs_info)

doc.add_heading('10.2 Jenkins Monitoring', 2)
jenkins_monitor = """Metrics Tracked:
  • Build success/failure rates
  • Build duration trends
  • Pipeline execution history
  • Artifact storage usage
  • Executor availability

Access: http://65.0.117.219:8080/monitoring

Features:
  • Email notifications on build failure
  • Build status badges
  • Trend graphs and statistics"""
doc.add_paragraph(jenkins_monitor)

doc.add_heading('10.3 SonarQube Quality Monitoring', 2)
sonar_monitor = """Metrics Dashboards:
  • Code quality trends over time
  • Coverage evolution
  • Bug discovery rate
  • Vulnerability tracking
  • Technical debt calculations

Alerts: Configured quality gates trigger alerts when standards degrade"""
doc.add_paragraph(sonar_monitor)

doc.add_page_break()

# Section 11: Security
doc.add_heading('11. Security & Best Practices', 1)

doc.add_heading('11.1 Network Security', 2)
doc.add_paragraph('[IMAGE PLACEHOLDER: Security Layers Diagram]')
network_sec = """Layers of Protection:

1. Security Groups (AWS Firewall)
   - Stateful filtering of network traffic
   - Whitelist allowed ports and sources
   - Restrict SSH access to known IPs

2. SSH Key-Based Authentication
   - Public-private key pairs instead of passwords
   - Private keys stored securely (not in Git)
   - SSH key: ~/.ssh/stayngo-deployer-key

3. Network Isolation
   - Default VPC provides basic isolation
   - Internal services use private IPs
   - Database behind Supabase (managed security)

HTTPS/TLS (Future Enhancement):
   - Use AWS Certificate Manager for SSL certificates
   - CloudFront CDN for HTTPS edge caching
   - Redirect HTTP to HTTPS"""
doc.add_paragraph(network_sec)

doc.add_heading('11.2 Secrets Management', 2)
secrets_mgmt = """Current Approach:
  • .env files on EC2 instance (not in Git repository)
  • Environment variables injected at runtime
  • Docker secrets for sensitive data

Stored Secrets:
  • Supabase URL and API key
  • Database credentials
  • JWT signing keys
  • Docker registry credentials

Best Practices:
  ✓ Use .env.example as template (committed to Git)
  ✓ Never commit actual .env files
  ✓ Limit file permissions (600)
  ✓ Rotate credentials regularly
  ✓ Use AWS Secrets Manager (recommended for production)"""
doc.add_paragraph(secrets_mgmt)

doc.add_heading('11.3 Database Security', 2)
db_sec = """Supabase PostgreSQL:
  • Encrypted at rest (AWS RDS encryption)
  • SSL/TLS in transit
  • Regular automated backups
  • Point-in-time recovery
  • Row-level security policies
  • Authentication role-based access

Application Security:
  • Prepared statements (SQL injection prevention)
  • Input validation and sanitization
  • Rate limiting on API endpoints"""
doc.add_paragraph(db_sec)

doc.add_heading('11.4 Container Security', 2)
container_sec = """Docker Best Practices:
  ✓ Use minimal base images (alpine, slim variants)
  ✓ Don't run containers as root
  ✓ Set resource limits (memory, CPU)
  ✓ Use read-only root filesystem where possible
  ✓ Scan images for vulnerabilities

Registry Security:
  • Private Docker Hub repositories (recommended)
  • Image signing and verification
  • Remove unused images regularly
  • Tag and version images properly"""
doc.add_paragraph(container_sec)

doc.add_page_break()

# Section 12: Deployment Pipeline
doc.add_heading('12. Deployment Pipeline Workflow', 1)

doc.add_heading('12.1 Complete Pipeline Flow', 2)
doc.add_paragraph('[IMAGE PLACEHOLDER: Complete CI/CD Pipeline Diagram]')

pipeline_flow = """STEP 1: Developer Commits Code to GitHub
  └─ Pushes changes to main branch
  └─ GitHub sends webhook to Jenkins (http://jenkins-ip:8080/github-webhook/)

STEP 2: Jenkins Webhook Triggered
  └─ Jenkins receives notification
  └─ Creates new build job in queue

STEP 3: Code Checkout
  └─ Jenkins clones repository: git clone github.com/Mayank08082004/StayNGo
  └─ Checks out the specific commit

STEP 4: SonarQube Analysis
  └─ Runs sonar-scanner on codebase
  └─ Analyzes for bugs, vulnerabilities, code smells
  └─ If quality gates fail → Pipeline STOPS

STEP 5: Docker Image Build
  ├─ Backend:
  │  └─ docker build -t mayankc1533262/stayngo-backend:latest ./backend --no-cache
  └─ Frontend:
     └─ docker build -t mayankc1533262/stayngo-frontend:latest ./frontend --no-cache

STEP 6: Docker Login & Push
  └─ docker login -u $DOCKER_USER -p $DOCKER_PASS
  └─ docker push mayankc1533262/stayngo-backend:latest
  └─ docker push mayankc1533262/stayngo-frontend:latest

STEP 7: SSH to EC2 Instance
  └─ ssh -i ~/.ssh/stayngo-deployer-key ubuntu@65.0.117.219
  └─ Establishes secure connection to production server

STEP 8: Deploy via Docker Compose
  └─ cd ~/stayngo
  └─ docker-compose pull (latest images from Docker Hub)
  └─ docker-compose up -d (start containers in background)

STEP 9: Health Checks
  └─ curl http://localhost:3000 (verify frontend)
  └─ curl http://localhost:5001/api (verify backend)

STEP 10: Pipeline Complete
  └─ Success notification sent (email, Slack, etc.)
  └─ Application updated to latest version"""
doc.add_paragraph(pipeline_flow)

doc.add_heading('12.2 Rollback Procedure', 2)
rollback_proc = """If deployment fails or issues occur:

1. SSH to EC2 instance
   ssh -i ~/.ssh/stayngo-deployer-key ubuntu@65.0.117.219

2. View running containers
   docker ps -a

3. Check recent Docker Compose history
   docker-compose logs -f stayngo-backend
   docker-compose logs -f stayngo-frontend

4. Rollback to previous version
   docker-compose pull (pulls current tag again)
   
   OR manually specify previous image version:
   Edit docker-compose.yml to use tagged version (e.g., :v1.2.0)
   docker-compose up -d

5. Verify services recovered
   curl http://localhost:3000
   curl http://localhost:5001/api"""
doc.add_paragraph(rollback_proc)

doc.add_page_break()

# Section 13: Troubleshooting
doc.add_heading('13. Troubleshooting & Maintenance', 1)

doc.add_heading('13.1 Common Issues & Solutions', 2)

doc.add_heading('Issue 1: Supabase Keys Not Found at Build Time', 3)
doc.add_paragraph('Problem: Next.js build fails with "NEXT_PUBLIC_* keys not available"')
doc.add_paragraph('Solution:', style='List Bullet')
doc.add_paragraph('1. Verify .env.production exists on build system', style='List Bullet 2')
doc.add_paragraph('2. Ensure Jenkins passes environment variables to Docker build', style='List Bullet 2')
doc.add_paragraph('3. Use --build-arg to pass secrets (NEXT_PUBLIC_* only)', style='List Bullet 2')
doc.add_paragraph('4. Review Jenkins logs: jenkins web console -> Build logs', style='List Bullet 2')

doc.add_heading('Issue 2: Backend Connection Refused', 3)
doc.add_paragraph('Problem: Frontend cannot reach backend (ECONNREFUSED 127.0.0.1:5001)')
doc.add_paragraph('Solution:', style='List Bullet')
doc.add_paragraph('1. Backend might not be running: docker ps | grep stayngo-backend', style='List Bullet 2')
doc.add_paragraph('2. Verify NEXT_PUBLIC_API_URL set correctly (not localhost inside Docker)', style='List Bullet 2')
doc.add_paragraph('3. Check security group allows port 5001', style='List Bullet 2')
doc.add_paragraph('4. Test directly: curl http://65.0.117.219:5001/api', style='List Bullet 2')

doc.add_heading('Issue 3: Gunicorn Worker Timeout', 3)
doc.add_paragraph('Problem: Backend freezes with Worker Timeout error')
doc.add_paragraph('Solution:', style='List Bullet')
doc.add_paragraph('1. Caused by slow client attacks (empty socket connections)', style='List Bullet 2')
doc.add_paragraph('2. Solution already configured: --workers 3 --threads 4', style='List Bullet 2')
doc.add_paragraph('3. Increase timeout if needed (currently 120s)', style='List Bullet 2')
doc.add_paragraph('4. Monitor: docker logs stayngo-backend', style='List Bullet 2')

doc.add_heading('Issue 4: Docker Image Pull Failed', 3)
doc.add_paragraph('Problem: docker pull mayankc1533262/stayngo-backend fails')
doc.add_paragraph('Solution:', style='List Bullet')
doc.add_paragraph('1. Check internet connectivity: ping 8.8.8.8', style='List Bullet 2')
doc.add_paragraph('2. Verify Docker daemon running: sudo systemctl status docker', style='List Bullet 2')
doc.add_paragraph('3. Login to Docker Hub: docker login', style='List Bullet 2')
doc.add_paragraph('4. Verify image exists: docker search username/image-name', style='List Bullet 2')

doc.add_heading('Issue 5: Terraform State Lock', 3)
doc.add_paragraph('Problem: "Error acquiring the state lock" when running terraform')
doc.add_paragraph('Solution:', style='List Bullet')
doc.add_paragraph('1. Another terraform process running: ps aux | grep terraform', style='List Bullet 2')
doc.add_paragraph('2. Force unlock: terraform force-unlock LOCK_ID', style='List Bullet 2')
doc.add_paragraph('3. Backup state file first!', style='List Bullet 2')
doc.add_paragraph('4. Re-run terraform apply', style='List Bullet 2')

doc.add_heading('13.2 Maintenance Tasks', 2)

doc.add_paragraph('Daily:', style='List Bullet')
doc.add_paragraph('• Monitor SonarQube quality gates', style='List Bullet 2')
doc.add_paragraph('• Check Jenkins build status', style='List Bullet 2')
doc.add_paragraph('• Review application logs for errors', style='List Bullet 2')

doc.add_paragraph('Weekly:', style='List Bullet')
doc.add_paragraph('• Review cloud costs and resource utilization', style='List Bullet 2')
doc.add_paragraph('• Update dependency packages (when appropriate)', style='List Bullet 2')
doc.add_paragraph('• Test backup/recovery procedures', style='List Bullet 2')

doc.add_paragraph('Monthly:', style='List Bullet')
doc.add_paragraph('• Rotate secrets and API keys', style='List Bullet 2')
doc.add_paragraph('• Review and update security group rules', style='List Bullet 2')
doc.add_paragraph('• Performance tuning based on metrics', style='List Bullet 2')
doc.add_paragraph('• Update documentation', style='List Bullet 2')

doc.add_paragraph('Quarterly:', style='List Bullet')
doc.add_paragraph('• Major dependency updates', style='List Bullet 2')
doc.add_paragraph('• Infrastructure scaling assessment', style='List Bullet 2')
doc.add_paragraph('• Security audit and penetration testing', style='List Bullet 2')
doc.add_paragraph('• Disaster recovery drill', style='List Bullet 2')

doc.add_heading('13.3 Useful Commands Reference', 2)

commands = """Docker Commands:
  docker ps                              # List running containers
  docker logs -f [container_id]          # View container logs
  docker exec -it [container_id] bash    # Enter container shell
  docker stats                           # Monitor resource usage
  docker-compose restart                 # Restart all services
  docker-compose down                    # Stop and remove containers

Jenkins Commands:
  curl http://localhost:8080/restart     # Restart Jenkins
  curl http://localhost:8080/api/json    # Get Jenkins status

System Commands:
  df -h                                  # Disk space usage
  free -h                                # Memory usage
  top                                    # Process monitoring
  netstat -tlnp                          # Open ports/services
  ssh -i ~/.ssh/key ubuntu@IP            # SSH to EC2

Terraform Commands:
  terraform fmt                          # Format code
  terraform validate                     # Validate syntax
  terraform plan -out=tfplan             # Save plan to file
  terraform apply tfplan                 # Apply saved plan
  terraform state list                   # List resources in state"""
doc.add_paragraph(commands)

doc.add_page_break()

# Appendix
doc.add_heading('Appendix: Architecture Checklist', 1)

doc.add_paragraph('Pre-Deployment Checklist:', style='List Bullet')
doc.add_paragraph('☐ GitHub repository configured with main branch protection', style='List Bullet 2')
doc.add_paragraph('☐ Jenkins server running and accessible', style='List Bullet 2')
doc.add_paragraph('☐ SonarQube server running and configured', style='List Bullet 2')
doc.add_paragraph('☐ Docker Hub account created and credentials configured', style='List Bullet 2')
doc.add_paragraph('☐ AWS account with appropriate permissions', style='List Bullet 2')
doc.add_paragraph('☐ SSH key pair generated and stored securely', style='List Bullet 2')
doc.add_paragraph('☐ .env files prepared (never commit to Git)', style='List Bullet 2')
doc.add_paragraph('☐ Terraform state file backed up', style='List Bullet 2')
doc.add_paragraph('☐ All required ports open in security groups', style='List Bullet 2')
doc.add_paragraph('☐ Domain name configured (if using DNS)', style='List Bullet 2')
doc.add_paragraph('☐ HTTPS certificates ready (if using HTTPS)', style='List Bullet 2')
doc.add_paragraph('☐ Monitoring and alerting configured', style='List Bullet 2')
doc.add_paragraph('☐ Backup strategy documented and tested', style='List Bullet 2')

doc.add_paragraph('Post-Deployment Verification:', style='List Bullet')
doc.add_paragraph('☐ Frontend accessible and responsive', style='List Bullet 2')
doc.add_paragraph('☐ Backend API responding correctly', style='List Bullet 2')
doc.add_paragraph('☐ Database connected and queryable', style='List Bullet 2')
doc.add_paragraph('☐ Jenkins pipeline executing successfully', style='List Bullet 2')
doc.add_paragraph('☐ SonarQube analyzing code correctly', style='List Bullet 2')
doc.add_paragraph('☐ Logs being collected and viewable', style='List Bullet 2')
doc.add_paragraph('☐ Performance metrics within acceptable range', style='List Bullet 2')
doc.add_paragraph('☐ Security group rules verified', style='List Bullet 2')
doc.add_paragraph('☐ Backups configured and tested', style='List Bullet 2')
doc.add_paragraph('☐ Team trained on operational procedures', style='List Bullet 2')

doc.add_heading('Contact & Support', 1)
support_info = """For questions or issues related to this DevOps infrastructure:
  • Primary Contact: DevOps Team Lead
  • Email: devops@stayngo.local
  • Slack Channel: #devops-support
  • Documentation: GitHub Wiki
  • Incident Response: Follow on-call rotation

Emergency Contacts:
  • AWS Account Owner: [Name/Contact]
  • Security Team: [Email]
  • Database Administrator: [Email]"""
doc.add_paragraph(support_info)

# Save final document
final_path = r'c:\Users\mayan\OneDrive\Desktop\Mayank\StayNGo\StayNGo_DevOps_Documentation.docx'
doc.save(final_path)
print(f'✓ Complete DevOps documentation created successfully!')
print(f'✓ File: {final_path}')
print(f'\n✓ Document includes:')
print('  • Comprehensive technology explanations')
print('  • 15+ image placeholders for diagrams')
print('  • Step-by-step pipeline workflows')
print('  • Configuration details and examples')
print('  • Troubleshooting guides for 5+ common issues')
print('  • Security best practices')
print('  • Maintenance procedures')
print('  • Useful command references')
print('  • Pre/Post-deployment checklists')
